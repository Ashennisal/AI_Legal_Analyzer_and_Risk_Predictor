import os
import re
import json
from typing import List, Optional
from docx import Document

# -------------------------
# 1) DOCX -> TEXT
# -------------------------
def read_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    lines: List[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    lines.append(t)

    return "\n".join(lines)


# -------------------------
# 2) Helpers
# -------------------------
ISO_DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")
HHMM_RE = re.compile(r"^\d{2}:\d{2}$")

def _coerce_time_hhmm(t: Optional[str], default_time: str) -> str:
    t = (t or "").strip()

    # "9:00" -> "09:00"
    if re.match(r"^\d{1,2}:\d{2}$", t):
        h, m = t.split(":")
        return f"{int(h):02d}:{int(m):02d}"

    if HHMM_RE.match(t):
        return t

    return default_time

def _coerce_date_iso(d: Optional[str]) -> Optional[str]:
    d = (d or "").strip()
    return d if ISO_DATE_RE.match(d) else None


def _strip_json_fences(s: str) -> str:
    """Remove ```json ... ``` or ``` ... ``` wrappers if model returns them."""
    s = s.strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.IGNORECASE)
        s = re.sub(r"\s*```$", "", s)
    return s.strip()


# -------------------------
# 3) Gemini Extractor
# -------------------------
def extract_events_with_gemini(
    text: str,
    default_time: str = "09:00",
    model: str = "gemini-2.5-flash",
) -> List[dict]:
    """
    Returns:
      [{"title":"...", "date":"YYYY-MM-DD", "time":"HH:MM"}, ...]
    """
    try:
        from google import genai
    except ImportError:
        raise RuntimeError(
            "google-genai library is not installed. "
            'Install with: pip install google-genai'
        )
    
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "GEMINI_API_KEY is not set.\n"
            'PowerShell:  $env:GEMINI_API_KEY="YOUR_KEY"\n'
            "Then:       python app.py"
        )

    client = genai.Client(
        api_key=api_key,
        http_options={"api_version": "v1"},
    )

    # reduce tokens / quota usage
    text = text[:12000]

    prompt = f"""
Return ONLY valid JSON. Do NOT use markdown. Do NOT wrap in ```.

Output format must be EXACTLY one of these:

Option A:
{{"events":[{{"title":"...","date":"YYYY-MM-DD","time":"HH:MM"}}]}}

Option B:
[{{"title":"...","date":"YYYY-MM-DD","time":"HH:MM"}}]

Rules:
- Extract explicit or strongly implied deadlines, due dates, renewals, expiries, meetings, appointments, hearings, follow-ups, admissions, tests.
- Only include events with real dates.
- If time missing, use "{default_time}".
- date must be YYYY-MM-DD
- time must be HH:MM (24-hour)

DOCUMENT:
{text}
""".strip()

    response = client.models.generate_content(
        model=model,
        contents=prompt,
    )

    raw = _strip_json_fences(response.text)

    try:
        data = json.loads(raw)
    except Exception:
        print("⚠ Model did not return valid JSON.")
        print("Raw response:\n", response.text)
        return []

    # Support both shapes
    if isinstance(data, list):
        events_raw = data
    else:
        events_raw = data.get("events", [])

    cleaned: List[dict] = []
    seen = set()

    for ev in events_raw:
        if not isinstance(ev, dict):
            continue

        iso_date = _coerce_date_iso(ev.get("date"))
        if not iso_date:
            continue

        hhmm = _coerce_time_hhmm(ev.get("time"), default_time)
        title = (ev.get("title") or "").strip() or "Important Deadline"

        normalized_title = title.lower().strip()
        key = (normalized_title, iso_date, hhmm)

        if key in seen:
            continue
        seen.add(key)

        cleaned.append({
        "title": title,
        "date": iso_date,
        "time": hhmm
        })
        
    return cleaned


def summarize_document_with_gemini(
    text: str,
    model: str = "gemini-2.5-flash",
) -> dict:
    """
    Returns {"technical": str, "layman": str, "actionable": str}.
    On failure or missing deps, returns three empty strings (caller may omit from API response).
    """
    empty = {"technical": "", "layman": "", "actionable": ""}
    if not (text or "").strip():
        return empty

    try:
        from google import genai
    except ImportError:
        return empty

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return empty

    client = genai.Client(
        api_key=api_key,
        http_options={"api_version": "v1"},
    )

    text = text[:12000]

    prompt = f"""
Return ONLY valid JSON. Do NOT use markdown. Do NOT wrap in ```.

Output format EXACTLY:
{{"technical":"...","layman":"...","actionable":"..."}}

Rules:
- "technical": concise summary for someone with legal training (key terms, structure, obligations).
- "layman": plain English for a non-lawyer; avoid jargon or explain it briefly.
- "actionable": concrete next steps, deadlines to watch, or decisions the reader should make.
- Each value must be a single JSON string (escape quotes and newlines as needed). Keep each under 1200 words.

DOCUMENT:
{text}
""".strip()

    try:
        response = client.models.generate_content(
            model=model,
            contents=prompt,
        )
        raw = _strip_json_fences(response.text or "")
        data = json.loads(raw)
    except Exception as e:
        print(f"⚠ summarize_document_with_gemini failed: {e}")
        return empty

    if not isinstance(data, dict):
        return empty

    out = {
        "technical": str(data.get("technical") or "").strip(),
        "layman": str(data.get("layman") or "").strip(),
        "actionable": str(data.get("actionable") or "").strip(),
    }
    return out
