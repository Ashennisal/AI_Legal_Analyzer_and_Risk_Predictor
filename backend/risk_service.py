import io
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Lazy NLTK/VADER so importing this module (and starting FastAPI) does not block on NLTK downloads.
_sent_tokenize = None
_analyzer = None


def _ensure_nlp():
    global _sent_tokenize, _analyzer
    if _analyzer is not None:
        return
    import nltk
    from nltk.tokenize import sent_tokenize
    from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer

    for _pkg in ("punkt", "punkt_tab"):
        try:
            nltk.data.find(f"tokenizers/{_pkg}")
        except LookupError:
            try:
                nltk.download(_pkg, quiet=True)
            except Exception:
                pass

    _sent_tokenize = sent_tokenize
    _analyzer = SentimentIntensityAnalyzer()


RISK_PHRASES = [
    "breach of contract", "financial loss", "penalty", "non compliance",
    "data breach", "lawsuit", "legal action", "violation", "terminated",
    "risk", "liability", "damages", "fraud", "unauthorized", "confidential"
]


def analyze_document_risk(text: str):
    """Analyzes text for risky sentences and phrases."""
    _ensure_nlp()
    risky_sentences = []
    risky_phrases_found = []

    sentences = _sent_tokenize(text.lower())

    for sent in sentences:
        for phrase in RISK_PHRASES:
            if phrase in sent:
                risky_phrases_found.append(phrase)
                risky_sentences.append(sent)

        score = _analyzer.polarity_scores(sent)
        if score["compound"] < -0.5:
            risky_sentences.append(sent)

    unique_sentences = list(set(risky_sentences))
    unique_phrases = list(set(risky_phrases_found))

    risk_score = min(100, (len(unique_sentences) * 10) + (len(unique_phrases) * 5))

    risk_level = "Low"
    if risk_score > 30:
        risk_level = "Medium"
    if risk_score > 70:
        risk_level = "High"

    return {
        "risk_level": risk_level,
        "risk_score": risk_score,
        "clauses_detected": len(unique_sentences),
        "risky_phrases": unique_phrases,
        "risky_sentences": unique_sentences
    }


def create_highlighted_docx(original_text: str, risky_sentences: list) -> io.BytesIO:
    """Creates a Word document with risky sentences highlighted in yellow."""
    doc = Document()
    para = doc.add_paragraph()

    sentences = original_text.split(".")
    risky_set = set([s.strip().lower() for s in risky_sentences])

    for s in sentences:
        clean_s = s.strip().lower()
        run = para.add_run(s + ". ")
        if clean_s in risky_set and clean_s != "":
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


HIGH_SIGNAL_PHRASES = [
    "breach of contract",
    "lawsuit",
    "legal action",
    "unlimited liability",
    "gross negligence",
    "indemnif",
]


def sentence_risk_level(sentence: str) -> str | None:
    """
    Per-sentence severity for UI highlighting: 'high', 'medium', 'low', or None.
    Uses phrase lists plus VADER compound (on original casing for sentiment).
    """
    _ensure_nlp()
    sl = sentence.lower()
    for phrase in HIGH_SIGNAL_PHRASES:
        if phrase in sl:
            return "high"
    score = _analyzer.polarity_scores(sentence)["compound"]
    if score < -0.52:
        return "high"
    for phrase in RISK_PHRASES:
        if phrase in sl:
            return "medium"
    if score < -0.38:
        return "medium"
    if score < -0.2:
        return "low"
    return None


def _find_sentence_span(text: str, sent: str, start: int) -> tuple[int, int] | None:
    for candidate in (sent, sent.strip()):
        if not candidate:
            continue
        pos = text.find(candidate, start)
        if pos != -1:
            return pos, pos + len(candidate)
    return None


def _merge_adjacent_same_level(segments: list[dict]) -> list[dict]:
    if not segments:
        return []
    merged = [{"text": segments[0]["text"], "level": segments[0]["level"]}]
    for seg in segments[1:]:
        if seg["level"] == merged[-1]["level"]:
            merged[-1]["text"] += seg["text"]
        else:
            merged.append({"text": seg["text"], "level": seg["level"]})
    return merged


def build_risk_segments(text: str) -> list[dict]:
    """
    Ordered segments covering the full document text for client-side highlighting.
    Each item: {"text": str, "level": None | 'low' | 'medium' | 'high'}.
    """
    _ensure_nlp()
    if not text:
        return []
    sents = _sent_tokenize(text)
    out: list[dict] = []
    i = 0
    n = len(text)
    for sent in sents:
        if not sent.strip():
            continue
        span = _find_sentence_span(text, sent, i)
        if span is None:
            continue
        pos, end = span
        if pos > i:
            out.append({"text": text[i:pos], "level": None})
        level = sentence_risk_level(sent)
        out.append({"text": text[pos:end], "level": level})
        i = end
    if i < n:
        out.append({"text": text[i:n], "level": None})
    return _merge_adjacent_same_level(out)


def detect_risks(text: str):
    """Returns (risky_sentences, risky_phrases) for main.analyze_uploaded_document."""
    result = analyze_document_risk(text)
    return result["risky_sentences"], result["risky_phrases"]


def extract_text_from_pdf(file_obj) -> str:
    """Read text from a PDF file-like object (binary mode)."""
    import PyPDF2

    pdf_reader = PyPDF2.PdfReader(file_obj)
    parts = []
    for page in pdf_reader.pages:
        t = page.extract_text()
        parts.append(t or "")
    return "\n".join(parts)
